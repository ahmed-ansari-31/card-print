import streamlit as st
from PIL import Image, ImageEnhance, ImageOps
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import os
import logging
from datetime import datetime

# === Setup ===
CARD_WIDTH_INCH = 3.375
CARD_HEIGHT_INCH = 2.125
PAGE_WIDTH, PAGE_HEIGHT = A4
PREVIEW_WIDTH = 600

log_dir = "logs"
output_dir = "output"
os.makedirs(log_dir, exist_ok=True)
os.makedirs(output_dir, exist_ok=True)

logging.basicConfig(
    filename=os.path.join(log_dir, "app.log"),
    filemode='a',
    format="%(asctime)s [%(levelname)s]: %(message)s",
    level=logging.INFO
)

# === Functions ===
def enhance_image(img, brightness=1.0, contrast=1.0, sharpness=1.0, grayscale=False,
                  rotate_angle=0, crop_values=(0, 0, 0, 0)):
    img = img.convert("RGB")
    if grayscale:
        img = ImageOps.grayscale(img).convert("RGB")
    img = img.rotate(rotate_angle, expand=True)

    width, height = img.size
    left_crop, top_crop, right_crop, bottom_crop = crop_values
    img = img.crop((
        left_crop,
        top_crop,
        width - right_crop,
        height - bottom_crop
    ))

    img = ImageEnhance.Brightness(img).enhance(brightness)
    img = ImageEnhance.Contrast(img).enhance(contrast)
    img = ImageEnhance.Sharpness(img).enhance(sharpness)
    return img


def resize_image(img):
    return img.resize((int(CARD_WIDTH_INCH * 300), int(CARD_HEIGHT_INCH * 300)))  # 300 DPI

def create_pdf(img, filename=None):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    card_width = CARD_WIDTH_INCH * inch
    card_height = CARD_HEIGHT_INCH * inch

    columns, rows = 2, 4
    h_spacing, v_spacing = 0.25 * inch, 0.25 * inch

    total_width = columns * card_width + (columns - 1) * h_spacing
    total_height = rows * card_height + (rows - 1) * v_spacing

    left_margin = (PAGE_WIDTH - total_width) / 2
    bottom_margin = (PAGE_HEIGHT - total_height) / 2

    x_positions = [left_margin + col * (card_width + h_spacing) for col in range(columns)]
    y_positions = [bottom_margin + row * (card_height + v_spacing) for row in reversed(range(rows))]

    preview = Image.new("RGB", (int(PAGE_WIDTH), int(PAGE_HEIGHT)), "white")
    for y in y_positions:
        for x in x_positions:
            c.drawInlineImage(img, x, y, width=card_width, height=card_height)
            preview.paste(img.resize((int(card_width), int(card_height))), (int(x), int(PAGE_HEIGHT - y - card_height)))

    c.showPage()
    c.save()
    buffer.seek(0)

    if filename:
        save_path = os.path.join(output_dir, filename)
        with open(save_path, "wb") as f:
            f.write(buffer.getbuffer())
        logging.info(f"PDF saved: {save_path}")

    return buffer, preview

def create_word(img, filename=None):
    doc = Document()
    for _ in range(4):
        table = doc.add_table(rows=1, cols=2)
        row_cells = table.rows[0].cells
        for cell in row_cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            image_io = BytesIO()
            img.save(image_io, format='PNG')
            image_io.seek(0)
            run.add_picture(image_io, width=Inches(CARD_WIDTH_INCH), height=Inches(CARD_HEIGHT_INCH))
    doc.add_page_break()
    word_io = BytesIO()
    doc.save(word_io)

    if filename:
        save_path = os.path.join(output_dir, filename)
        with open(save_path, "wb") as f:
            f.write(word_io.getbuffer())
        logging.info(f"Word saved: {save_path}")

    word_io.seek(0)
    return word_io

# === Session state ===
if "image_uploaded" not in st.session_state:
    st.session_state.image_uploaded = False
if "preview_ready" not in st.session_state:
    st.session_state.preview_ready = False

# === App UI ===
st.title("🪪 ID Card Duplicator & Print Generator")

# === Upload ===
if not st.session_state.image_uploaded:
    uploaded_file1 = st.file_uploader("Upload an ID card image", type=["jpg", "jpeg", "png"])
    if uploaded_file1:
        st.session_state.original_image = Image.open(uploaded_file1).convert("RGB")
        st.session_state.image_uploaded = True
        st.success("✅ Image uploaded. You can now edit it below.")

# === Real-time editor ===
if st.session_state.image_uploaded:
    st.subheader("🎛 Real-Time Image Editor (Single Card)")
    grayscale = st.toggle("Grayscale")
    brightness = st.slider("Brightness", 0.5, 2.0, 1.0, 0.05)
    contrast = st.slider("Contrast", 0.5, 2.0, 1.0, 0.05)
    sharpness = st.slider("Sharpness", 0.5, 2.0, 1.0, 0.05)

    edited_preview = enhance_image(
        st.session_state.original_image,
        brightness,
        contrast,
        sharpness,
        grayscale
    )
    st.image(edited_preview, caption="🔍 Single Card Preview", width=PREVIEW_WIDTH)

    if st.button("🧱 Apply to Layout & Generate PDF/Word"):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        resized_img = resize_image(edited_preview)

        pdf_buf, layout_preview = create_pdf(resized_img, f"id_cards_{timestamp}.pdf")
        word_buf = create_word(resized_img, f"id_cards_{timestamp}.docx")

        st.session_state.preview_ready = True
        st.session_state.preview_img = layout_preview
        st.session_state.pdf_data = pdf_buf
        st.session_state.word_data = word_buf

# === Final layout preview + download ===
if st.session_state.get("preview_ready", False):
    st.subheader("🖼 Final Layout (8 cards on A4)")
    st.image(st.session_state.preview_img, use_column_width=True)

    st.download_button("📄 Download PDF", data=st.session_state.pdf_data, file_name="ID_Cards.pdf", mime="application/pdf")
    st.download_button("📝 Download Word", data=st.session_state.word_data, file_name="ID_Cards.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if st.button("🖨 Print from Browser"):
        st.markdown('<script>window.print()</script>', unsafe_allow_html=True)

    if st.button("🔁 Reset"):
        st.session_state.clear()  # Clear all session state
        st.rerun() 
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image

def save_image_as_word(images, positions, page_width_inch, page_height_inch, img_width_inch, img_height_inch, dpi=300, auto_rotate=False):
    """
    images: list of PIL Images (already cropped/enhanced)
    positions: list of (x_inch, y_inch, rotate90) for each image (top-left corner, in inches, and rotation flag)
    page_width_inch, page_height_inch: page size
    img_width_inch, img_height_inch: image size (before rotation)
    auto_rotate: if True, rotate image 90deg if it fits better
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(page_width_inch)
    section.page_height = Inches(page_height_inch)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0)
    for img, (x_inch, y_inch, rotate90) in zip(images, positions):
        image_io = BytesIO()
        if auto_rotate and rotate90:
            img = img.rotate(90, expand=True)
            w, h = img_height_inch, img_width_inch
        else:
            w, h = img_width_inch, img_height_inch
        img.save(image_io, format='PNG')
        image_io.seek(0)
        # Add a 1x1 table to position the image
        table = doc.add_table(rows=1, cols=1)
        table.allow_autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(w)
        cell.height = Inches(h)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_io, width=Inches(w), height=Inches(h))
        # Set table position by adding empty paragraphs (simulate y offset)
        for _ in range(int(y_inch * dpi / 72)):
            doc.add_paragraph("")
        # Set left margin for x offset (simulate x offset)
        table.style = 'Table Grid'
        table.autofit = False
        # Note: Word does not support absolute positioning, so this is a best-effort
    word_io = BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io 